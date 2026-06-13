from __future__ import annotations

import argparse
import io
import mimetypes
import os
import re
import smtplib
import ssl
import sys
from contextlib import redirect_stderr
from dataclasses import dataclass
from datetime import datetime
from email.message import EmailMessage
from email.utils import formataddr
from pathlib import Path
from typing import Any

import pandas as pd


INPUT_ROOT = Path("main_input")
RESUME_DIR_NAME = "resume"
EXCEL_DIR_NAME = "excel"
OUTPUT_DIR = Path("output")
DEFAULT_OUTPUT_FILE = OUTPUT_DIR / "generated_cold_emails.xlsx"
DEFAULT_ROLES = "Data Scientist, Data Engineer, Data Analyst"
DEFAULT_MODEL = "gemini-3-flash-preview"
GENERATED_EMAIL_COLUMN = "Generated Email"
EMAIL_SUBJECT_COLUMN = "Email Subject"
EMAIL_BODY_COLUMN = "Email Body"
EMAIL_SEND_STATUS_COLUMN = "Email Send Status"
EMAIL_SENT_AT_COLUMN = "Email Sent At"
GENERIC_HR_NAME = "Hiring Team"
GENERIC_PROFILE_SKILLS = [
    "Python",
    "SQL",
    "Machine Learning",
    "Data Analysis",
    "Data Visualization",
    "Statistics",
    "Pandas",
    "NumPy",
    "Scikit-learn",
    "Power BI",
    "Excel",
    "ETL",
]

RESUME_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xls", ".csv"}

COMPANY_COLUMN_ALIASES = {
    "company",
    "company name",
    "company_name",
    "organization",
    "organisation",
    "companyname",
    "employer",
}
HR_NAME_COLUMN_ALIASES = {
    "hr",
    "hr name",
    "hr_name",
    "recruiter",
    "recruiter name",
    "contact",
    "contact name",
    "hiring manager",
    "name",
}
HR_EMAIL_COLUMN_ALIASES = {
    "hremail",
    "hr email",
    "hr_email",
    "email",
    "email id",
    "email address",
    "recruiter email",
    "contact email",
    "mail",
}
GENERATED_EMAIL_COLUMN_ALIASES = {
    "generated email",
    "generated_email",
    "email content",
    "email body",
    "mail body",
    "cold email",
}
GENERIC_EMAIL_LOCAL_PARTS = {
    "admin",
    "apply",
    "career",
    "careers",
    "contact",
    "cv",
    "email",
    "enquiry",
    "hello",
    "hr",
    "hiring",
    "info",
    "job",
    "jobs",
    "mail",
    "noreply",
    "no",
    "people",
    "recruiter",
    "recruiters",
    "recruiting",
    "recruitment",
    "reply",
    "resume",
    "resumes",
    "support",
    "talent",
    "team",
    "work",
}
EMAIL_PATTERN = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)

KNOWN_SKILLS = [
    "Python",
    "R",
    "SQL",
    "Excel",
    "Power BI",
    "Tableau",
    "Machine Learning",
    "Deep Learning",
    "NLP",
    "Computer Vision",
    "Statistics",
    "Data Analysis",
    "Data Visualization",
    "Predictive Modeling",
    "Regression",
    "Classification",
    "Clustering",
    "A/B Testing",
    "ETL",
    "Data Engineering",
    "Data Warehousing",
    "Spark",
    "PySpark",
    "Hadoop",
    "Airflow",
    "dbt",
    "AWS",
    "Azure",
    "GCP",
    "Docker",
    "Git",
    "Linux",
    "Pandas",
    "NumPy",
    "Scikit-learn",
    "TensorFlow",
    "PyTorch",
    "Keras",
    "Matplotlib",
    "Seaborn",
    "Plotly",
    "BigQuery",
    "Snowflake",
    "PostgreSQL",
    "MySQL",
    "MongoDB",
    "NoSQL",
    "APIs",
    "Flask",
    "FastAPI",
    "MLOps",
    "Model Deployment",
    "Feature Engineering",
    "Generative AI",
    "LLM",
]


@dataclass
class ResumeProfile:
    text: str
    candidate_name: str
    skills: list[str]
    highlights: list[str]


@dataclass
class PreparedEmail:
    row_index: int
    recipient: str
    hr_name: str
    subject: str
    body: str


@dataclass
class SMTPSettings:
    host: str
    port: int
    username: str
    password: str
    from_email: str
    from_name: str
    reply_to: str
    use_tls: bool
    use_ssl: bool


@dataclass
class EmailSendSummary:
    prepared: int = 0
    sent: int = 0
    failed: int = 0
    skipped: int = 0
    dry_run: bool = False


def load_env_file(path: Path) -> None:
    """Load simple KEY=VALUE pairs from .env without adding another dependency."""
    if not path.exists():
        return

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip().lstrip("\ufeff")
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def clean_cell(value: Any) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def normalize_column_name(name: str) -> str:
    normalized = re.sub(r"[_\-]+", " ", str(name).strip().lower())
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized


def compact_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def find_files(directory: Path, extensions: set[str]) -> list[Path]:
    if not directory.exists():
        return []
    return sorted(
        path
        for path in directory.rglob("*")
        if path.is_file() and path.suffix.lower() in extensions
    )


def extract_pdf_text(path: Path, allow_gemini_pdf: bool = True) -> str:
    errors = []
    try:
        import fitz

        fitz.TOOLS.mupdf_display_errors(False)
        fitz.TOOLS.mupdf_display_warnings(False)
        with redirect_stderr(io.StringIO()):
            document = fitz.open(path)
        with document:
            with redirect_stderr(io.StringIO()):
                text = "\n".join(page.get_text("text") for page in document)
            if text.strip():
                return text
            errors.append(
                f"PyMuPDF: no extractable text found across {document.page_count} pages"
            )
    except Exception as exc:
        errors.append(f"PyMuPDF: {exc}")

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        errors.append(f"pypdf import: {exc}")
    else:
        try:
            with redirect_stderr(io.StringIO()):
                reader = PdfReader(str(path), strict=False)
                pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
            if text.strip():
                return text
            errors.append(f"pypdf: no extractable text found across {len(reader.pages)} pages")
        except Exception as exc:
            errors.append(f"pypdf: {exc}")

    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if allow_gemini_pdf and api_key:
        try:
            return extract_pdf_text_with_gemini(path)
        except Exception as exc:
            errors.append(f"Gemini PDF extraction: {exc}")

    detail = "; ".join(errors)
    raise RuntimeError(
        "Could not read PDF resume. Export the resume as DOCX/TXT, or set GEMINI_API_KEY "
        "so Gemini can read the PDF directly. "
        f"Details: {detail}"
    )


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install python-docx to read DOCX resumes: pip install python-docx") from exc

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_cells.append(cell.text)
    return "\n".join(paragraphs + table_cells)


def extract_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def extract_resume_text(path: Path, allow_gemini_pdf: bool = True) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path, allow_gemini_pdf=allow_gemini_pdf)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix in {".txt", ".md"}:
        return extract_text_file(path)
    raise ValueError(f"Unsupported resume format: {path.name}")


def extract_candidate_name(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:8]:
        if "@" in line or re.search(r"\d{5,}", line):
            continue
        if len(line.split()) <= 5 and re.search(r"[A-Za-z]", line):
            return line
    return "Candidate"


def extract_skill_section(text: str) -> list[str]:
    lines = [line.strip() for line in text.splitlines()]
    collected: list[str] = []
    capture = False

    heading_pattern = re.compile(
        r"^(technical\s+skills|skills|tools|technologies|core\s+competencies)\b",
        re.IGNORECASE,
    )
    stop_pattern = re.compile(
        r"^(experience|work\s+experience|projects|education|certifications|summary|profile|achievements)\b",
        re.IGNORECASE,
    )

    for line in lines:
        if heading_pattern.match(line):
            capture = True
            remainder = re.sub(heading_pattern, "", line).strip(" :-|")
            if remainder:
                collected.append(remainder)
            continue
        if capture and stop_pattern.match(line):
            break
        if capture and line:
            collected.append(line)
        if capture and len(collected) >= 10:
            break

    raw = " ".join(collected)
    if not raw:
        return []

    parts = re.split(r"[,;|/]|(?:\s+-\s+)|(?:\s+and\s+)", raw)
    skills = []
    for part in parts:
        skill = re.sub(r"\s+", " ", part.strip(" .:-"))
        if 1 <= len(skill.split()) <= 5 and 1 < len(skill) <= 40:
            skills.append(skill)
    return dedupe(skills)


def extract_known_skills(text: str) -> list[str]:
    found: list[str] = []
    lower_text = text.lower()
    for skill in KNOWN_SKILLS:
        pattern = r"(?<![A-Za-z0-9+#.])" + re.escape(skill.lower()) + r"(?![A-Za-z0-9+#.])"
        if re.search(pattern, lower_text):
            found.append(skill)
    return found


def extract_highlights(text: str, max_items: int = 5) -> list[str]:
    important_terms = re.compile(
        r"\b(project|experience|intern|developed|built|created|implemented|analyzed|"
        r"dashboard|model|machine learning|data|analytics|prediction|classification|"
        r"visualization|pipeline|automation|research|certification)\b",
        re.IGNORECASE,
    )
    lines = []
    for line in text.splitlines():
        line = re.sub(r"^[\-*•\d.)\s]+", "", line.strip())
        if 30 <= len(line) <= 220 and important_terms.search(line):
            lines.append(line)
    return dedupe(lines)[:max_items]


def dedupe(values: list[str]) -> list[str]:
    seen = set()
    output = []
    for value in values:
        normalized = re.sub(r"\s+", " ", value.strip()).lower()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        output.append(value.strip())
    return output


def candidate_name_from_filename(path: Path) -> str:
    stem = re.sub(r"[_\-]+", " ", path.stem)
    stem = re.sub(r"\b(resume|cv|profile)\b", "", stem, flags=re.IGNORECASE)
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem.title() if stem else "Candidate"


def build_generic_resume_profile(resume_dir: Path, reason: str) -> ResumeProfile:
    resume_files = find_files(resume_dir, RESUME_EXTENSIONS)
    candidate_name = (
        candidate_name_from_filename(resume_files[0]) if resume_files else "Candidate"
    )
    text = (
        f"Resume text could not be extracted. Reason: {reason}\n"
        f"Candidate name inferred from resume filename: {candidate_name}"
    )
    return ResumeProfile(
        text=text,
        candidate_name=candidate_name,
        skills=GENERIC_PROFILE_SKILLS,
        highlights=[
            "Interested in data scientist, data engineering, and analyst opportunities.",
            "Able to contribute to data-driven problem solving, reporting, and analytical workflows.",
        ],
    )


def build_resume_profile(
    resume_dir: Path,
    allow_gemini_pdf: bool = True,
    allow_generic_resume: bool = False,
) -> ResumeProfile:
    resume_files = find_files(resume_dir, RESUME_EXTENSIONS)
    if not resume_files:
        raise FileNotFoundError(
            f"No resume file found in {resume_dir}. Add a .pdf, .docx, .txt, or .md resume."
        )

    sections = []
    for path in resume_files:
        try:
            text = compact_text(extract_resume_text(path, allow_gemini_pdf=allow_gemini_pdf))
        except Exception as exc:
            if allow_generic_resume:
                print(
                    "Warning: Could not extract resume text; using generic resume fallback. "
                    f"Reason: {exc}",
                    file=sys.stderr,
                )
                return build_generic_resume_profile(resume_dir, str(exc))
            raise
        if text:
            sections.append(f"Resume file: {path.name}\n{text}")

    combined_text = compact_text("\n\n".join(sections))
    if not combined_text:
        if allow_generic_resume:
            return build_generic_resume_profile(
                resume_dir,
                "Resume files were found, but no readable text could be extracted.",
            )
        raise ValueError("Resume files were found, but no readable text could be extracted.")

    skills = dedupe(extract_skill_section(combined_text) + extract_known_skills(combined_text))
    highlights = extract_highlights(combined_text)
    return ResumeProfile(
        text=combined_text,
        candidate_name=extract_candidate_name(combined_text),
        skills=skills[:35],
        highlights=highlights,
    )


def read_spreadsheet(path: Path) -> pd.DataFrame:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        raw = pd.read_csv(path, header=None)
    else:
        raw = pd.read_excel(path, header=None)

    raw = raw.dropna(how="all").dropna(axis=1, how="all")
    if raw.empty:
        return raw

    first_row = [clean_cell(value) for value in raw.iloc[0].tolist()]
    aliases = COMPANY_COLUMN_ALIASES | HR_NAME_COLUMN_ALIASES | HR_EMAIL_COLUMN_ALIASES
    looks_like_header = any(normalize_column_name(value) in aliases for value in first_row)

    if looks_like_header:
        columns = []
        used_columns = set()
        for index, value in enumerate(first_row, start=1):
            column = value or f"Column {index}"
            if column in used_columns:
                column = f"{column} {index}"
            used_columns.add(column)
            columns.append(column)
        frame = raw.iloc[1:].copy()
        frame.columns = columns
        return frame

    frame = raw.copy()
    frame.columns = [f"Column {index}" for index in range(1, len(frame.columns) + 1)]
    return frame


def load_company_rows(excel_dir: Path) -> pd.DataFrame:
    files = find_files(excel_dir, SPREADSHEET_EXTENSIONS)
    if not files:
        raise FileNotFoundError(
            f"No company/HR spreadsheet found in {excel_dir}. Add a .xlsx, .xls, or .csv file."
        )

    frames = []
    for path in files:
        frame = read_spreadsheet(path)
        frame["Source File"] = path.name
        frames.append(frame)

    combined = pd.concat(frames, ignore_index=True)
    combined = combined.dropna(how="all")
    if combined.empty:
        raise ValueError("Company/HR spreadsheets were found, but they do not contain rows.")
    return combined


def find_column(columns: list[str], aliases: set[str]) -> str | None:
    normalized_map = {normalize_column_name(column): column for column in columns}
    for alias in aliases:
        if alias in normalized_map:
            return normalized_map[alias]
    for normalized, original in normalized_map.items():
        if normalized in aliases:
            return original
    return None


def extract_email_address(value: Any) -> str:
    match = EMAIL_PATTERN.search(clean_cell(value))
    return match.group(0).strip() if match else ""


def contains_email(value: Any) -> bool:
    return bool(extract_email_address(value))


def derive_hr_name_from_email(value: Any) -> str:
    email = extract_email_address(value).lower()
    if not email:
        return GENERIC_HR_NAME

    local_part = email.split("@", 1)[0].split("+", 1)[0]
    local_part = re.sub(r"\d+", " ", local_part)
    tokens = [
        token.lower()
        for token in re.split(r"[^A-Za-z]+", local_part)
        if token.strip()
    ]
    personal_tokens = [
        token
        for token in tokens
        if token not in GENERIC_EMAIL_LOCAL_PARTS and len(token) > 1
    ]
    if not personal_tokens:
        return GENERIC_HR_NAME
    return " ".join(token.capitalize() for token in personal_tokens[:2])


def infer_email_column(frame: pd.DataFrame) -> str | None:
    best_column = None
    best_count = 0
    for column in frame.columns:
        count = sum(contains_email(value) for value in frame[column])
        if count > best_count:
            best_column = column
            best_count = count
    return best_column if best_count else None


def is_phone_like(value: str) -> bool:
    compact = re.sub(r"\D", "", value)
    return len(compact) >= 7 and len(compact) >= max(1, len(value.replace(" ", "")) - 3)


def company_from_email(value: Any) -> str:
    match = EMAIL_PATTERN.search(clean_cell(value))
    if not match:
        return ""
    domain = match.group(0).split("@", 1)[1].lower()
    parts = domain.split(".")
    if len(parts) >= 2 and parts[-2] in {"co", "com", "org", "net"}:
        name = parts[-3] if len(parts) >= 3 else parts[0]
    else:
        name = parts[-2] if len(parts) >= 2 else parts[0]
    name = re.sub(r"[^a-z0-9]+", " ", name).strip()
    return name.title() if name else ""


def text_score(series: pd.Series) -> int:
    score = 0
    for value in series:
        text = clean_cell(value)
        if not text or contains_email(text) or is_phone_like(text):
            continue
        if re.search(r"[A-Za-z]", text):
            score += 1
    return score


def infer_company_column(frame: pd.DataFrame, email_column: str | None) -> str | None:
    candidates = [
        column
        for column in frame.columns
        if column != email_column and normalize_column_name(column) != "source file"
    ]
    if not candidates:
        return None

    generic_columns = [
        column
        for column in candidates
        if re.fullmatch(r"Column \d+", str(column)) and text_score(frame[column])
    ]
    if generic_columns:
        generic_columns.sort(key=lambda column: int(str(column).split()[1]))
        return generic_columns[0]

    scored = [(text_score(frame[column]), column) for column in candidates]
    scored.sort(key=lambda item: (-item[0], str(item[1])))
    return scored[0][1] if scored[0][0] else candidates[0]


def infer_hr_name_column(
    frame: pd.DataFrame,
    company_column: str | None,
    email_column: str | None,
) -> str | None:
    candidates = [
        column
        for column in frame.columns
        if column not in {company_column, email_column}
        and normalize_column_name(column) != "source file"
    ]
    if not candidates:
        return None
    scored = [(text_score(frame[column]), column) for column in candidates]
    scored.sort(key=lambda item: (-item[0], str(item[1])))
    return scored[0][1] if scored[0][0] else None


def row_to_context(row: pd.Series) -> dict[str, str]:
    context = {}
    for key, value in row.items():
        clean_value = clean_cell(value)
        if clean_value:
            context[str(key)] = clean_value
    return context


def fallback_row_company(row: pd.Series, primary_column: str, hr_name_column: str | None) -> str:
    company_name = clean_cell(row[primary_column])
    if company_name:
        return company_name

    if hr_name_column:
        hr_detail = clean_cell(row[hr_name_column])
        if hr_detail and not contains_email(hr_detail) and not is_phone_like(hr_detail):
            return hr_detail

    for column, value in row.items():
        if normalize_column_name(column) == "source file":
            continue
        text = clean_cell(value)
        if text and not contains_email(text) and not is_phone_like(text):
            return text

    for value in row.values:
        company = company_from_email(value)
        if company:
            return company
    return ""


def truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rsplit(" ", 1)[0] + "..."


def extract_pdf_text_with_gemini(path: Path) -> str:
    from google import genai

    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    model = os.environ.get("GEMINI_MODEL", DEFAULT_MODEL)
    client = genai.Client(api_key=api_key)
    uploaded_file = client.files.upload(file=path)
    prompt = (
        "Extract the complete readable text from this resume PDF. Preserve names, "
        "contact details, section headings, skills, projects, education, experience, "
        "certifications, and tools. Output plain text only."
    )
    response = client.models.generate_content(
        model=model,
        contents=[prompt, uploaded_file],
    )
    text = (response.text or "").strip()
    if not text:
        raise RuntimeError("Gemini returned no text from the resume PDF.")
    return text


def can_use_gemini(no_llm: bool) -> bool:
    if no_llm:
        return False
    if not (os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")):
        return False
    try:
        from google import genai  # noqa: F401
    except ImportError:
        return False
    return True


def generate_email_with_gemini(
    *,
    company_name: str,
    hr_name: str,
    roles: str,
    resume: ResumeProfile,
    row_context: dict[str, str],
    model: str,
) -> str:
    from google import genai
    from google.genai import types

    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    client = genai.Client(api_key=api_key)

    system_prompt = (
        "You write polished, honest cold emails for job opportunities. "
        "Use only the supplied resume and company/HR details. Do not invent degrees, employers, "
        "metrics, projects, or experience. Keep the tone confident, concise, and human."
    )
    user_prompt = f"""
Create one cold email for an HR/recruiting contact.

Required outcome:
- The candidate is interested in {roles} roles.
- The candidate is looking for relevant opportunities at the company.
- Make it specific to the company and HR contact when details are available.
- Mention the candidate's strongest relevant skills and why they are worth considering.
- Include a subject line.
- Use a greeting with the HR name if available; otherwise use "Hiring Team".
- Keep the email between 150 and 220 words.
- Do not include placeholders.

Company/HR details:
{row_context}

Candidate name:
{resume.candidate_name}

Extracted skills:
{", ".join(resume.skills) if resume.skills else "Not explicitly extracted"}

Resume highlights:
{chr(10).join("- " + item for item in resume.highlights) if resume.highlights else "No separate highlights extracted"}

Resume text excerpt:
{truncate(resume.text, 9000)}
""".strip()

    response = client.models.generate_content(
        model=model,
        contents=user_prompt,
        config=types.GenerateContentConfig(system_instruction=system_prompt),
    )
    email = (response.text or "").strip()
    if not email:
        raise RuntimeError("Gemini returned an empty response.")
    return email


def generate_fallback_email(
    *,
    company_name: str,
    hr_name: str,
    roles: str,
    resume: ResumeProfile,
    row_context: dict[str, str],
) -> str:
    greeting_name = hr_name if hr_name else "Hiring Team"
    company_display = company_name if company_name else "your company"
    skills = ", ".join(resume.skills[:12]) if resume.skills else "data analysis, problem solving, and analytical thinking"
    highlights = " ".join(resume.highlights[:2])
    if highlights:
        highlights_sentence = f" A few relevant highlights from my background include: {highlights}"
    else:
        highlights_sentence = ""

    extra_context = ""
    for key in ("Industry", "Domain", "Location", "Job Title", "Role", "Company Details"):
        value = row_context.get(key)
        if value:
            extra_context = f" I noticed your work around {value}, which makes the opportunity especially interesting to me."
            break

    return f"""Subject: Interest in {roles} opportunities at {company_display}

Dear {greeting_name},

I hope you are doing well. I am reaching out to express my interest in {roles} opportunities at {company_display}.{extra_context}

My resume reflects a strong foundation in {skills}.{highlights_sentence} I enjoy working on data-driven problems, building useful analytical solutions, and translating business questions into clear insights.

I would be grateful if you could consider my profile for any suitable openings in data science, data engineering, or analyst roles. I would also be happy to share more details about my projects, experience, and how I can contribute to your team.

Thank you for your time and consideration.

Regards,
{resume.candidate_name}"""


def generate_email(
    *,
    company_name: str,
    hr_name: str,
    roles: str,
    resume: ResumeProfile,
    row_context: dict[str, str],
    model: str,
    use_llm: bool,
) -> tuple[str, str]:
    if use_llm:
        try:
            email = generate_email_with_gemini(
                company_name=company_name,
                hr_name=hr_name,
                roles=roles,
                resume=resume,
                row_context=row_context,
                model=model,
            )
            return email, "gemini"
        except Exception as exc:
            print(
                f"Warning: Gemini generation failed for {company_name or 'unknown company'}; "
                f"using fallback template. Reason: {exc}",
                file=sys.stderr,
            )

    email = generate_fallback_email(
        company_name=company_name,
        hr_name=hr_name,
        roles=roles,
        resume=resume,
        row_context=row_context,
    )
    return email, "fallback"


def find_generated_email_column(columns: list[str]) -> str | None:
    for column in columns:
        if normalize_column_name(column) == normalize_column_name(GENERATED_EMAIL_COLUMN):
            return column
    return find_column(columns, GENERATED_EMAIL_COLUMN_ALIASES)


def load_generated_output(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"Generated email output not found: {path}")
    if path.suffix.lower() == ".csv":
        return pd.read_csv(path)
    return pd.read_excel(path)


def save_generated_output(frame: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.suffix.lower() == ".csv":
        frame.to_csv(path, index=False)
    else:
        frame.to_excel(path, index=False)


def split_generated_email(generated_email: Any) -> tuple[str, str]:
    text = clean_cell(generated_email).replace("\r\n", "\n").replace("\r", "\n")
    text = text.strip()
    if not text:
        return "", ""

    lines = text.split("\n")
    subject = ""
    body_lines = lines
    for index, line in enumerate(lines[:8]):
        match = re.match(
            r"^\s*\*{0,2}\s*subject\s*[:\-\*]*\s*(.+?)\s*$",
            line.strip(),
            re.IGNORECASE,
        )
        if match:
            subject = re.sub(r"\s+", " ", match.group(1).strip().strip("*")).strip()
            body_lines = lines[index + 1 :]
            break

    body = "\n".join(body_lines).strip()
    return subject or "Interest in data opportunities", body or text


def sanitize_hr_name(name: Any) -> str:
    clean_name = re.sub(r"[\r\n,]+", " ", clean_cell(name))
    clean_name = re.sub(r"\s+", " ", clean_name).strip()
    return clean_name or GENERIC_HR_NAME


def personalize_email_body(body: str, hr_name: str) -> str:
    clean_name = sanitize_hr_name(hr_name)
    text = body.strip()
    greeting = f"Dear {clean_name},"
    greeting_pattern = re.compile(r"(?im)^\s*(?:dear|hello|hi)\s+[^\n,]+,?")
    if greeting_pattern.search(text):
        return greeting_pattern.sub(greeting, text, count=1).strip()
    return f"{greeting}\n\n{text}".strip()


def parse_bool_env(value: str | None, default: bool) -> bool:
    if value is None or not str(value).strip():
        return default
    return str(value).strip().lower() in {"1", "true", "yes", "y", "on"}


def load_smtp_settings() -> SMTPSettings:
    host = clean_cell(os.environ.get("SMTP_HOST"))
    if not host:
        raise ValueError(
            "SMTP_HOST is required to send emails. Add SMTP settings to .env first."
        )

    use_ssl = parse_bool_env(os.environ.get("SMTP_USE_SSL"), False)
    port = int(os.environ.get("SMTP_PORT") or (465 if use_ssl else 587))
    use_ssl = parse_bool_env(os.environ.get("SMTP_USE_SSL"), port == 465)
    use_tls = parse_bool_env(os.environ.get("SMTP_USE_TLS"), not use_ssl)
    username = clean_cell(os.environ.get("SMTP_USERNAME"))
    password = clean_cell(os.environ.get("SMTP_PASSWORD"))
    from_email = clean_cell(os.environ.get("SMTP_FROM_EMAIL")) or username
    from_name = clean_cell(os.environ.get("SMTP_FROM_NAME"))
    reply_to = clean_cell(os.environ.get("SMTP_REPLY_TO"))

    if not from_email:
        raise ValueError("SMTP_FROM_EMAIL or SMTP_USERNAME is required to send emails.")
    if username and not password:
        raise ValueError("SMTP_PASSWORD is required when SMTP_USERNAME is set.")

    return SMTPSettings(
        host=host,
        port=port,
        username=username,
        password=password,
        from_email=from_email,
        from_name=from_name,
        reply_to=reply_to,
        use_tls=use_tls,
        use_ssl=use_ssl,
    )


def add_attachments(message: EmailMessage, attachments: list[Path]) -> None:
    for path in attachments:
        mime_type, _ = mimetypes.guess_type(path.name)
        if not mime_type:
            mime_type = "application/octet-stream"
        maintype, subtype = mime_type.split("/", 1)
        message.add_attachment(
            path.read_bytes(),
            maintype=maintype,
            subtype=subtype,
            filename=path.name,
        )


def send_one_email(
    prepared_email: PreparedEmail,
    settings: SMTPSettings,
    attachments: list[Path],
) -> None:
    message = EmailMessage()
    message["From"] = formataddr((settings.from_name, settings.from_email))
    message["To"] = prepared_email.recipient
    message["Subject"] = prepared_email.subject
    if settings.reply_to:
        message["Reply-To"] = settings.reply_to
    message.set_content(prepared_email.body)
    add_attachments(message, attachments)

    context = ssl.create_default_context()
    if settings.use_ssl:
        with smtplib.SMTP_SSL(settings.host, settings.port, context=context) as server:
            if settings.username:
                server.login(settings.username, settings.password)
            server.send_message(message)
        return

    with smtplib.SMTP(settings.host, settings.port) as server:
        if settings.use_tls:
            server.starttls(context=context)
        if settings.username:
            server.login(settings.username, settings.password)
        server.send_message(message)


def collect_attachments(
    *,
    input_root: Path,
    attach_resume: bool,
    attachment_paths: list[Path],
) -> list[Path]:
    attachments: list[Path] = []
    if attach_resume:
        resume_dir = input_root / RESUME_DIR_NAME
        attachments.extend(
            path
            for path in find_files(resume_dir, RESUME_EXTENSIONS)
            if path.name.lower() != "readme.md"
        )
    attachments.extend(attachment_paths)

    missing = [str(path) for path in attachments if not path.is_file()]
    if missing:
        raise FileNotFoundError("Attachment file not found: " + ", ".join(missing))
    return attachments


def prepare_emails_from_output(
    frame: pd.DataFrame,
    send_limit: int | None,
) -> tuple[list[PreparedEmail], int]:
    columns = list(frame.columns)
    hr_email_column = find_column(columns, HR_EMAIL_COLUMN_ALIASES) or infer_email_column(frame)
    generated_email_column = find_generated_email_column(columns)
    hr_name_column = find_column(columns, HR_NAME_COLUMN_ALIASES) or "HR Name"

    if not hr_email_column:
        raise ValueError("Could not find the HREMAIL column in the generated output.")
    if not generated_email_column:
        raise ValueError("Could not find the Generated Email column in the generated output.")
    if hr_name_column not in frame.columns:
        frame[hr_name_column] = ""

    for column in (
        EMAIL_SUBJECT_COLUMN,
        EMAIL_BODY_COLUMN,
        EMAIL_SEND_STATUS_COLUMN,
        EMAIL_SENT_AT_COLUMN,
    ):
        if column not in frame.columns:
            frame[column] = ""
        frame[column] = frame[column].astype("object")

    prepared: list[PreparedEmail] = []
    skipped = 0
    effective_limit = send_limit if send_limit and send_limit > 0 else None

    for row_index, row in frame.iterrows():
        if effective_limit is not None and len(prepared) >= effective_limit:
            break

        recipient = extract_email_address(row[hr_email_column])
        generated_email = clean_cell(row[generated_email_column])
        if not recipient:
            frame.at[row_index, EMAIL_SEND_STATUS_COLUMN] = "Skipped: missing HREMAIL"
            skipped += 1
            continue
        if not generated_email:
            frame.at[row_index, EMAIL_SEND_STATUS_COLUMN] = "Skipped: missing Generated Email"
            skipped += 1
            continue

        hr_name = derive_hr_name_from_email(recipient)
        subject, body = split_generated_email(generated_email)
        body = personalize_email_body(body, hr_name)
        frame.at[row_index, hr_name_column] = hr_name
        frame.at[row_index, EMAIL_SUBJECT_COLUMN] = subject
        frame.at[row_index, EMAIL_BODY_COLUMN] = body
        prepared.append(
            PreparedEmail(
                row_index=row_index,
                recipient=recipient,
                hr_name=hr_name,
                subject=subject,
                body=body,
            )
        )

    return prepared, skipped


def send_generated_emails(
    frame: pd.DataFrame,
    *,
    dry_run: bool,
    send_limit: int | None,
    attachments: list[Path],
) -> EmailSendSummary:
    prepared_emails, skipped = prepare_emails_from_output(frame, send_limit)
    summary = EmailSendSummary(
        prepared=len(prepared_emails),
        skipped=skipped,
        dry_run=dry_run,
    )
    if not prepared_emails:
        return summary

    settings = None if dry_run else load_smtp_settings()
    if attachments:
        names = ", ".join(path.name for path in attachments)
        print(f"Attachments: {names}")

    for position, prepared_email in enumerate(prepared_emails, start=1):
        if dry_run:
            frame.at[prepared_email.row_index, EMAIL_SEND_STATUS_COLUMN] = (
                "Dry run: ready"
            )
            print(
                f"Dry run {position}/{len(prepared_emails)}: "
                f"{prepared_email.recipient} | {prepared_email.subject}"
            )
            continue

        try:
            assert settings is not None
            send_one_email(prepared_email, settings, attachments)
        except Exception as exc:
            frame.at[prepared_email.row_index, EMAIL_SEND_STATUS_COLUMN] = f"Failed: {exc}"
            summary.failed += 1
            print(
                f"Failed {position}/{len(prepared_emails)}: "
                f"{prepared_email.recipient} ({exc})",
                file=sys.stderr,
            )
            continue

        frame.at[prepared_email.row_index, EMAIL_SEND_STATUS_COLUMN] = "Sent"
        frame.at[prepared_email.row_index, EMAIL_SENT_AT_COLUMN] = datetime.now().isoformat(
            timespec="seconds"
        )
        summary.sent += 1
        print(f"Sent {position}/{len(prepared_emails)}: {prepared_email.recipient}")

    return summary


def create_output(
    *,
    input_root: Path,
    output_path: Path,
    roles: str,
    model: str,
    no_llm: bool,
    allow_generic_resume: bool,
) -> pd.DataFrame:
    resume_dir = input_root / RESUME_DIR_NAME
    excel_dir = input_root / EXCEL_DIR_NAME

    resume = build_resume_profile(
        resume_dir,
        allow_gemini_pdf=not no_llm,
        allow_generic_resume=allow_generic_resume,
    )
    company_rows = load_company_rows(excel_dir)

    company_column = find_column(list(company_rows.columns), COMPANY_COLUMN_ALIASES)
    hr_name_column = find_column(list(company_rows.columns), HR_NAME_COLUMN_ALIASES)
    hr_email_column = find_column(list(company_rows.columns), HR_EMAIL_COLUMN_ALIASES)

    if not hr_email_column:
        hr_email_column = infer_email_column(company_rows)
        if hr_email_column:
            print(f"Inferred HR email column: {hr_email_column}")
    if not company_column:
        company_column = infer_company_column(company_rows, hr_email_column)
        if company_column:
            print(f"Inferred company column: {company_column}")
    if not hr_name_column:
        hr_name_column = infer_hr_name_column(company_rows, company_column, hr_email_column)
        if hr_name_column:
            print(f"Inferred HR/contact detail column: {hr_name_column}")

    if not company_column:
        raise ValueError(
            "Could not find the company column. Use a column like 'Company name' or 'Company'."
        )
    if not hr_email_column:
        raise ValueError(
            "Could not find the HR email column. Use a column like 'HREMAIL', 'HR Email', or 'Email'."
        )

    use_llm = can_use_gemini(no_llm)
    if use_llm:
        print(f"Using Gemini model: {model}")
    else:
        print(
            "Using fallback template mode. Add GEMINI_API_KEY and install google-genai "
            "for Gemini mode."
        )

    output_rows = []
    for index, row in company_rows.iterrows():
        row_context = row_to_context(row)
        company_name = fallback_row_company(row, company_column, hr_name_column)
        hr_email = clean_cell(row[hr_email_column])
        hr_name = derive_hr_name_from_email(hr_email)
        if hr_name_column:
            row_context[str(hr_name_column)] = hr_name
        row_context["Resolved HR Name"] = hr_name

        email, mode = generate_email(
            company_name=company_name,
            hr_name=hr_name,
            roles=roles,
            resume=resume,
            row_context=row_context,
            model=model,
            use_llm=use_llm,
        )
        output_rows.append(
            {
                "Company name": company_name,
                "HR Name": hr_name,
                "HREMAIL": hr_email,
                "Generated Email": email,
            }
        )
        print(
            f"Generated {index + 1}/{len(company_rows)}: "
            f"{company_name or 'Unknown company'} ({mode})"
        )

    output = pd.DataFrame(output_rows)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output.to_excel(output_path, index=False)
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate company- and HR-specific cold emails from a resume and HR spreadsheet."
    )
    parser.add_argument(
        "--input-root",
        type=Path,
        default=INPUT_ROOT,
        help="Main input folder containing resume/ and excel/ subfolders.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT_FILE,
        help="Path for the generated Excel output.",
    )
    parser.add_argument(
        "--roles",
        default=DEFAULT_ROLES,
        help="Comma-separated roles to target in the generated emails.",
    )
    parser.add_argument(
        "--model",
        default=os.environ.get("GEMINI_MODEL", DEFAULT_MODEL),
        help="Gemini model to use when LLM mode is available.",
    )
    parser.add_argument(
        "--no-llm",
        action="store_true",
        help="Disable Gemini generation and use the local fallback template.",
    )
    parser.add_argument(
        "--allow-generic-resume",
        action="store_true",
        help="Generate emails with a generic data profile if resume text cannot be extracted.",
    )
    parser.add_argument(
        "--send-emails",
        action="store_true",
        help="Send emails after generating the output workbook.",
    )
    parser.add_argument(
        "--send-existing-output",
        action="store_true",
        help="Skip generation and send emails from the existing --output workbook.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Preview recipients and subjects without connecting to SMTP.",
    )
    parser.add_argument(
        "--send-limit",
        type=int,
        default=None,
        help="Maximum number of prepared emails to send or preview.",
    )
    parser.add_argument(
        "--attach-resume",
        action="store_true",
        help="Attach resume files found in main_input/resume/ to each email.",
    )
    parser.add_argument(
        "--attachment",
        type=Path,
        action="append",
        default=[],
        help="Extra file to attach to each email. Repeat this option for multiple files.",
    )
    return parser.parse_args()


def main() -> int:
    load_env_file(Path(".env"))
    args = parse_args()

    try:
        if args.send_existing_output:
            output = load_generated_output(args.output)
            print(f"Loaded {len(output)} generated emails from: {args.output}")
        else:
            output = create_output(
                input_root=args.input_root,
                output_path=args.output,
                roles=args.roles,
                model=args.model,
                no_llm=args.no_llm,
                allow_generic_resume=args.allow_generic_resume,
            )

        should_send = args.send_emails or args.send_existing_output or args.dry_run
        if should_send:
            attachments = collect_attachments(
                input_root=args.input_root,
                attach_resume=args.attach_resume,
                attachment_paths=args.attachment,
            )
            summary = send_generated_emails(
                output,
                dry_run=args.dry_run,
                send_limit=args.send_limit,
                attachments=attachments,
            )
            if not args.dry_run:
                save_generated_output(output, args.output)
                print(f"Updated send status in: {args.output}")
            print(
                "Email summary: "
                f"prepared={summary.prepared}, sent={summary.sent}, "
                f"failed={summary.failed}, skipped={summary.skipped}"
            )
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    if args.send_existing_output:
        print(f"\nDone. Processed generated emails from: {args.output}")
    else:
        print(f"\nDone. Wrote {len(output)} generated emails to: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
